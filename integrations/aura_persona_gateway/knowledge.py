"""Knowledge base (RAG) storage, chunking, embedding and retrieval.

FastGPT-style pipeline: kb -> document -> chunks with float32 embeddings in
sqlite. Retrieval is brute-force cosine similarity (embeddings are L2
normalized on write, so cosine == dot product). Pure python on purpose: the
container is stdlib-only besides websockets/pypdf/python-docx, and KB sizes
here are small (thousands of chunks).
"""
from __future__ import annotations

import json
import math
import re
import sqlite3
import time
import uuid
from array import array
from contextlib import closing
from io import BytesIO
from pathlib import Path
from threading import Lock
from typing import Any

from dataclasses import dataclass
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 80
EMBED_BATCH_SIZE = 32
EMBED_MAX_ATTEMPTS = 4
EMBED_RETRY_BASE_SECONDS = 1.5
SUPPORTED_FILE_TYPES = ("txt", "md", "pdf", "docx")
_SENTENCE_BREAKS = "。！？；!?;"


@dataclass(frozen=True)
class EmbeddingConfig:
    base_url: str = "https://api.jina.ai/v1"
    api_key: str = ""
    model: str = "jina-embeddings-v3"
    timeout_seconds: float = 30.0


class EmbeddingClient:
    """OpenAI-compatible /embeddings client.

    Works with Jina today; the same protocol is used by StepFun step_plan and
    other OpenAI-compatible endpoints, so swapping providers is config-only.
    The jina-only ``task`` parameter is attached just for jina-* models.
    """

    def __init__(self, config: EmbeddingConfig) -> None:
        self.config = config

    def embed(self, texts: list[str], *, task: str = "retrieval.passage") -> list[list[float]]:
        clean = [str(item or "") for item in texts]
        if not clean:
            return []
        base_url = str(self.config.base_url or "").strip().rstrip("/")
        if not base_url:
            raise ValueError("embedding base_url is not configured")
        api_key = str(self.config.api_key or "").strip()
        if not api_key:
            raise ValueError("embedding api_key is not configured")
        model = str(self.config.model or "").strip()
        if not model:
            raise ValueError("embedding model is not configured")
        body: dict[str, Any] = {"model": model, "input": clean}
        if model.startswith("jina-") and task:
            body["task"] = task
        req = Request(
            f"{base_url}/embeddings",
            data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
                # Jina 的 Cloudflare 会用 error 1010 拦掉 Python-urllib 默认 UA，必须带正常 UA。
                "User-Agent": "Mozilla/5.0 (compatible; aura-lily-kb/1.0)",
            },
        )
        timeout = max(1.0, float(self.config.timeout_seconds or 30.0))
        payload: dict[str, Any] | None = None
        # Jina 免费档并发上限很低（2），批量上传时容易撞 429，指数退避重试几次。
        for attempt in range(EMBED_MAX_ATTEMPTS):
            try:
                with urlopen(req, timeout=timeout) as res:
                    payload = json.loads(res.read().decode("utf-8"))
                break
            except HTTPError as exc:
                detail = ""
                try:
                    detail = exc.read().decode("utf-8", errors="replace")[:300]
                except Exception:
                    detail = ""
                if exc.code == 429 and attempt < EMBED_MAX_ATTEMPTS - 1:
                    time.sleep(EMBED_RETRY_BASE_SECONDS * (2 ** attempt))
                    continue
                raise RuntimeError(f"embedding HTTP {exc.code}: {detail}") from exc
            except (OSError, URLError, json.JSONDecodeError) as exc:
                raise RuntimeError(f"embedding request failed: {exc.__class__.__name__}") from exc
        if payload is None:
            raise RuntimeError("embedding request failed: no response")
        rows = payload.get("data")
        if not isinstance(rows, list) or len(rows) != len(clean):
            raise RuntimeError("embedding response shape mismatch")
        ordered = sorted(rows, key=lambda item: int(item.get("index") or 0))
        vectors: list[list[float]] = []
        for item in ordered:
            vector = item.get("embedding")
            if not isinstance(vector, list) or not vector:
                raise RuntimeError("embedding response missing vector")
            vectors.append([float(value) for value in vector])
        return vectors

    def embed_query(self, text: str) -> list[float]:
        vectors = self.embed([text], task="retrieval.query")
        return vectors[0]


def chunk_text(
    text: str,
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[str]:
    """FastGPT-style splitting: paragraphs -> lines -> sentences -> hard cut."""
    chunk_size = max(50, int(chunk_size or DEFAULT_CHUNK_SIZE))
    overlap = max(0, min(int(overlap or 0), chunk_size // 2))
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return []
    pieces: list[str] = []
    for paragraph in re.split(r"\n{2,}", normalized):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if len(paragraph) <= chunk_size:
            pieces.append(paragraph)
            continue
        pieces.extend(_split_long_text(paragraph, chunk_size))
    chunks: list[str] = []
    current = ""
    for piece in pieces:
        candidate = f"{current}\n{piece}" if current else piece
        if len(candidate) <= chunk_size:
            current = candidate
            continue
        if current:
            chunks.append(current)
            tail = current[-overlap:] if overlap else ""
            current = f"{tail}\n{piece}" if tail else piece
            if len(current) > chunk_size:
                chunks.extend(_split_long_text(current, chunk_size))
                current = ""
        else:
            chunks.extend(_split_long_text(piece, chunk_size))
    if current:
        chunks.append(current)
    return [chunk.strip() for chunk in chunks if chunk.strip()]


def _split_long_text(text: str, chunk_size: int) -> list[str]:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    sentences: list[str] = []
    for line in lines:
        if len(line) <= chunk_size:
            sentences.append(line)
            continue
        sentences.extend(_split_sentences(line, chunk_size))
    merged: list[str] = []
    current = ""
    for sentence in sentences:
        candidate = f"{current}{sentence}" if current else sentence
        if len(candidate) <= chunk_size:
            current = candidate
            continue
        if current:
            merged.append(current)
        current = sentence
    if current:
        merged.append(current)
    return merged


def _split_sentences(text: str, chunk_size: int) -> list[str]:
    sentences: list[str] = []
    buffer = ""
    for char in text:
        buffer += char
        if char in _SENTENCE_BREAKS and buffer.strip():
            sentences.append(buffer)
            buffer = ""
    if buffer.strip():
        sentences.append(buffer)
    result: list[str] = []
    for sentence in sentences:
        while len(sentence) > chunk_size:
            result.append(sentence[:chunk_size])
            sentence = sentence[chunk_size:]
        if sentence:
            result.append(sentence)
    return result


def extract_text(filename: str, raw: bytes) -> str:
    file_type = file_type_of(filename)
    if file_type in {"txt", "md"}:
        try:
            return raw.decode("utf-8-sig")
        except UnicodeDecodeError:
            return raw.decode("utf-8", errors="replace")
    if file_type == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf 未安装：请重建镜像（requirements.txt 已包含 pypdf）") from exc
        reader = PdfReader(BytesIO(raw))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip())
    if file_type == "docx":
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("python-docx 未安装：请重建镜像（requirements.txt 已包含 python-docx）") from exc
        document = docx.Document(BytesIO(raw))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    raise ValueError(f"unsupported file type: {filename}")


def file_type_of(filename: str) -> str:
    suffix = Path(str(filename or "")).suffix.lower().lstrip(".")
    return suffix if suffix in SUPPORTED_FILE_TYPES else ""


def default_kb_db_path(persona_home: str | Path) -> Path:
    return Path(persona_home).expanduser() / "knowledge" / "kb.sqlite3"


def default_kb_files_dir(persona_home: str | Path) -> Path:
    return Path(persona_home).expanduser() / "knowledge" / "files"


class KnowledgeStore:
    _schema_lock = Lock()
    _schema_ready_paths: set[str] = set()
    _retryable_sqlite_errors = ("disk i/o error", "database is locked", "database is busy", "locked")

    def __init__(self, db_path: str | Path) -> None:
        self.db_path = Path(db_path).expanduser()

    # ------------------------------------------------------------------ kb
    def list_kbs(self) -> list[dict[str, Any]]:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> list[dict[str, Any]]:
            rows = conn.execute(
                """
                SELECT kb.id, kb.name, kb.created_at, kb.updated_at,
                       COUNT(DISTINCT document.id) AS doc_count,
                       COUNT(chunk.id) AS chunk_count
                  FROM kb
                  LEFT JOIN document ON document.kb_id = kb.id
                  LEFT JOIN chunk ON chunk.kb_id = kb.id
                 GROUP BY kb.id
                 ORDER BY kb.created_at ASC
                """
            ).fetchall()
            return [dict(row) for row in rows]

        return self._execute_retryable(op)

    def create_kb(self, name: str) -> dict[str, Any]:
        self._ensure_schema()
        clean_name = str(name or "").strip()
        if not clean_name:
            raise ValueError("知识库名称不能为空")
        kb_id = f"kb-{uuid.uuid4().hex[:12]}"
        now = int(time.time())

        def op(conn: sqlite3.Connection) -> None:
            conn.execute(
                "INSERT INTO kb (id, name, created_at, updated_at) VALUES (?, ?, ?, ?)",
                (kb_id, clean_name, now, now),
            )
            conn.commit()

        self._execute_retryable(op)
        return {"id": kb_id, "name": clean_name, "created_at": now, "updated_at": now,
                "doc_count": 0, "chunk_count": 0}

    def get_kb(self, kb_id: str) -> dict[str, Any] | None:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> dict[str, Any] | None:
            row = conn.execute("SELECT * FROM kb WHERE id=? LIMIT 1", (str(kb_id or ""),)).fetchone()
            return dict(row) if row is not None else None

        return self._execute_retryable(op)

    def delete_kb(self, kb_id: str) -> bool:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> bool:
            cur = conn.execute("DELETE FROM kb WHERE id=?", (str(kb_id or ""),))
            conn.commit()
            return cur.rowcount > 0

        return self._execute_retryable(op)

    # ------------------------------------------------------------ document
    def list_documents(self, kb_id: str) -> list[dict[str, Any]]:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> list[dict[str, Any]]:
            rows = conn.execute(
                """
                SELECT id, kb_id, filename, file_type, char_count, chunk_count,
                       status, error, created_at, updated_at
                  FROM document
                 WHERE kb_id=?
                 ORDER BY created_at ASC
                """,
                (str(kb_id or ""),),
            ).fetchall()
            return [dict(row) for row in rows]

        return self._execute_retryable(op)

    def get_document(self, doc_id: str) -> dict[str, Any] | None:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> dict[str, Any] | None:
            row = conn.execute(
                "SELECT * FROM document WHERE id=? LIMIT 1", (str(doc_id or ""),)
            ).fetchone()
            return dict(row) if row is not None else None

        return self._execute_retryable(op)

    def create_document(self, kb_id: str, *, filename: str, file_type: str) -> dict[str, Any]:
        self._ensure_schema()
        doc_id = f"doc-{uuid.uuid4().hex[:12]}"
        now = int(time.time())

        def op(conn: sqlite3.Connection) -> None:
            conn.execute(
                """
                INSERT INTO document
                (id, kb_id, filename, file_type, char_count, chunk_count, status, error, created_at, updated_at)
                VALUES (?, ?, ?, ?, 0, 0, 'pending', '', ?, ?)
                """,
                (doc_id, str(kb_id or ""), str(filename or ""), str(file_type or ""), now, now),
            )
            conn.commit()

        self._execute_retryable(op)
        return {"id": doc_id, "kb_id": kb_id, "filename": filename, "file_type": file_type,
                "char_count": 0, "chunk_count": 0, "status": "pending", "error": "",
                "created_at": now, "updated_at": now}

    def set_document_status(
        self,
        doc_id: str,
        *,
        status: str,
        error: str = "",
        char_count: int | None = None,
        chunk_count: int | None = None,
    ) -> None:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> None:
            fields = ["status=?", "error=?", "updated_at=?"]
            values: list[Any] = [str(status or ""), str(error or "")[:500], int(time.time())]
            if char_count is not None:
                fields.append("char_count=?")
                values.append(int(char_count))
            if chunk_count is not None:
                fields.append("chunk_count=?")
                values.append(int(chunk_count))
            values.append(str(doc_id or ""))
            conn.execute(f"UPDATE document SET {', '.join(fields)} WHERE id=?", values)
            conn.commit()

        self._execute_retryable(op)

    def delete_document(self, doc_id: str) -> bool:
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> bool:
            cur = conn.execute("DELETE FROM document WHERE id=?", (str(doc_id or ""),))
            conn.commit()
            return cur.rowcount > 0

        return self._execute_retryable(op)

    # --------------------------------------------------------------- chunk
    def replace_chunks(
        self,
        doc_id: str,
        kb_id: str,
        chunks: list[str],
        embeddings: list[list[float]],
    ) -> int:
        if len(chunks) != len(embeddings):
            raise ValueError("chunks/embeddings length mismatch")
        self._ensure_schema()

        def op(conn: sqlite3.Connection) -> int:
            conn.execute("DELETE FROM chunk WHERE doc_id=?", (str(doc_id or ""),))
            count = 0
            for seq, (content, vector) in enumerate(zip(chunks, embeddings)):
                normalized = _l2_normalize(vector)
                blob = array("f", normalized).tobytes()
                conn.execute(
                    "INSERT INTO chunk (doc_id, kb_id, seq, content, embedding, dim) VALUES (?, ?, ?, ?, ?, ?)",
                    (str(doc_id or ""), str(kb_id or ""), seq, str(content or ""), blob, len(normalized)),
                )
                count += 1
            conn.commit()
            return count

        return self._execute_retryable(op)

    def search(
        self,
        kb_id: str,
        query_vec: list[float],
        *,
        top_k: int = 5,
        score_threshold: float = 0.45,
    ) -> list[dict[str, Any]]:
        self._ensure_schema()
        query = _l2_normalize(query_vec)
        if not query:
            return []
        top_k = max(1, min(20, int(top_k or 5)))
        threshold = max(0.0, min(1.0, float(score_threshold or 0.0)))

        def op(conn: sqlite3.Connection) -> list[dict[str, Any]]:
            rows = conn.execute(
                """
                SELECT chunk.id, chunk.doc_id, chunk.content, chunk.embedding, chunk.dim,
                       document.filename
                  FROM chunk
                  LEFT JOIN document ON document.id = chunk.doc_id
                 WHERE chunk.kb_id=?
                """,
                (str(kb_id or ""),),
            ).fetchall()
            hits: list[dict[str, Any]] = []
            for row in rows:
                vec = array("f")
                vec.frombytes(row["embedding"])
                if len(vec) != len(query):
                    continue
                score = sum(a * b for a, b in zip(query, vec))
                if score < threshold:
                    continue
                hits.append(
                    {
                        "chunk_id": int(row["id"]),
                        "doc_id": str(row["doc_id"]),
                        "filename": str(row["filename"] or ""),
                        "content": str(row["content"] or ""),
                        "score": float(score),
                    }
                )
            hits.sort(key=lambda item: item["score"], reverse=True)
            return hits[:top_k]

        return self._execute_retryable(op)

    # ------------------------------------------------------------ internal
    def _ensure_schema(self) -> None:
        key = str(self.db_path)
        if key in self._schema_ready_paths:
            return
        with self._schema_lock:
            if key in self._schema_ready_paths:
                return
            self.db_path.parent.mkdir(parents=True, exist_ok=True)
            with self._connection() as conn:
                conn.executescript(
                    """
                    CREATE TABLE IF NOT EXISTS kb (
                        id          TEXT PRIMARY KEY,
                        name        TEXT NOT NULL,
                        created_at  INTEGER NOT NULL,
                        updated_at  INTEGER NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS document (
                        id          TEXT PRIMARY KEY,
                        kb_id       TEXT NOT NULL REFERENCES kb(id) ON DELETE CASCADE,
                        filename    TEXT NOT NULL,
                        file_type   TEXT NOT NULL,
                        char_count  INTEGER NOT NULL DEFAULT 0,
                        chunk_count INTEGER NOT NULL DEFAULT 0,
                        status      TEXT NOT NULL DEFAULT 'pending',
                        error       TEXT NOT NULL DEFAULT '',
                        created_at  INTEGER NOT NULL,
                        updated_at  INTEGER NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS chunk (
                        id          INTEGER PRIMARY KEY AUTOINCREMENT,
                        doc_id      TEXT NOT NULL REFERENCES document(id) ON DELETE CASCADE,
                        kb_id       TEXT NOT NULL,
                        seq         INTEGER NOT NULL,
                        content     TEXT NOT NULL,
                        embedding   BLOB NOT NULL,
                        dim         INTEGER NOT NULL
                    );
                    CREATE INDEX IF NOT EXISTS idx_chunk_kb ON chunk(kb_id);
                    CREATE INDEX IF NOT EXISTS idx_chunk_doc ON chunk(doc_id);
                    CREATE INDEX IF NOT EXISTS idx_doc_kb ON document(kb_id);
                    """
                )
                conn.commit()
            self._schema_ready_paths.add(key)

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path, timeout=15)
        conn.execute("PRAGMA busy_timeout=15000")
        conn.execute("PRAGMA foreign_keys=ON")
        conn.row_factory = sqlite3.Row
        return conn

    def _connection(self):
        return closing(self._connect())

    def _execute_retryable(self, operation: Any) -> Any:
        last_error: sqlite3.OperationalError | None = None
        for attempt in range(3):
            try:
                with self._connection() as conn:
                    return operation(conn)
            except sqlite3.OperationalError as exc:
                last_error = exc
                if attempt >= 2 or not self._is_retryable_sqlite_error(exc):
                    raise
                time.sleep(0.05 * (attempt + 1))
        if last_error:
            raise last_error
        raise sqlite3.OperationalError("sqlite operation failed")

    def _is_retryable_sqlite_error(self, exc: sqlite3.OperationalError) -> bool:
        message = str(exc).lower()
        return any(item in message for item in self._retryable_sqlite_errors)


def _l2_normalize(vector: list[float] | array) -> list[float]:
    values = [float(item) for item in vector]
    norm = math.sqrt(sum(item * item for item in values))
    if norm <= 0:
        return []
    return [item / norm for item in values]


def process_document(
    store: KnowledgeStore,
    embedder: EmbeddingClient,
    *,
    doc_id: str,
    kb_id: str,
    filename: str,
    raw: bytes,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> dict[str, Any]:
    """Extract -> chunk -> embed -> store. Synchronous; run in a thread."""
    store.set_document_status(doc_id, status="processing")
    try:
        text = extract_text(filename, raw)
        if not text.strip():
            raise ValueError("文档内容为空或无法提取文本")
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
        if not chunks:
            raise ValueError("文档分块结果为空")
        embeddings: list[list[float]] = []
        for start in range(0, len(chunks), EMBED_BATCH_SIZE):
            batch = chunks[start : start + EMBED_BATCH_SIZE]
            embeddings.extend(embedder.embed(batch, task="retrieval.passage"))
        stored = store.replace_chunks(doc_id, kb_id, chunks, embeddings)
        store.set_document_status(
            doc_id, status="ready", char_count=len(text), chunk_count=stored
        )
        return {"ok": True, "doc_id": doc_id, "chunk_count": stored, "char_count": len(text)}
    except Exception as exc:  # noqa: BLE001 - status must always be recorded
        store.set_document_status(doc_id, status="failed", error=str(exc))
        return {"ok": False, "doc_id": doc_id, "error": str(exc)}


def kb_search(
    store: KnowledgeStore,
    embedder: EmbeddingClient,
    *,
    kb_id: str,
    query: str,
    top_k: int,
    score_threshold: float,
) -> list[dict[str, Any]]:
    """Embed the query and search. Raises on embedding API failure."""
    clean_query = str(query or "").strip()
    if not clean_query:
        return []
    query_vec = embedder.embed_query(clean_query)
    return store.search(kb_id, query_vec, top_k=top_k, score_threshold=score_threshold)
